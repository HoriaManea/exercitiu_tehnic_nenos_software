from docx import Document
from pathlib import Path
import re


def load_document(path: str) -> Document:
    if not Path(path).exists():
        raise FileNotFoundError(f"File not found: {path}")
    return Document(path)

def extract_table_paragraphs(table):
    paragraphs = []

    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)

            for inner_table in cell.tables:
                paragraphs.extend(extract_table_paragraphs(inner_table))                

    return placeholders
    
def find_placeholders(document: Document):
    placeholders = []

    pattern = re.compile(r"[_\.]{3,}")  

    for i, paragraph in enumerate(document.paragraphs):
        matches = pattern.findall(paragraph.text)
        if matches:
            placeholders.append({
                "type": "paragraph",
                "index": i,
                "text": paragraph.text,
                "matches": matches
            })

    return placeholders


if __name__ == "__main__":
    doc_path = "output_forms.docx"
    document = load_document(doc_path)

    placeholders = find_placeholders(document)

    print("\nFound placeholders:\n")
    for p in placeholders:
        print(p)