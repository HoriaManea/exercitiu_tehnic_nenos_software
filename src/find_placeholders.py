import re
from extract_table_paragraphs import extract_table_paragraphs
from docx import Document

def find_placeholders(document: Document):
    placeholders = []

    pattern = re.compile(r"[_\.]{3,}")  

    all_paragraphs = []

    all_paragraphs.extend(document.paragraphs)

    for table in document.tables:
        all_paragraphs.extend(extract_table_paragraphs(table))

    for section in document.sections:
        all_paragraphs.extend(section.header.paragraphs)
        all_paragraphs.extend(section.footer.paragraphs)

        for table in section.header.tables:
            all_paragraphs.extend(extract_table_paragraphs(table))

        for table in section.footer.tables:
            all_paragraphs.extend(extract_table_paragraphs(table))
        


    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text
            matches = list(pattern.finditer(run_text))

            for match in matches:
                placeholders.append({
                    "paragraph_obj": paragraph,
                    "run_obj": run,
                    "placeholder": match.group(),
                    "start_in_run": match.start(),
                    "end_in_run": match.end(),
                    "context": paragraph.text,
                    "runs_count": len(paragraph.runs)
                })

    return placeholders

