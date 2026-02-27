import re
from extract_table_paragraphs import extract_table_paragraphs
from docx import Document

def find_placeholders(document: Document):
    placeholders = []

    pattern = re.compile(r"[_\.]{3,}")  

    all_paragraphs = []

    all_paragraphs.extend(document.paragraphs)

    for table in document.tables:
        all_paragraphs.extend(extract_table_paragraphs(table))

    for section in document.sections:
        all_paragraphs.extend(section.header.paragraphs)
        all_paragraphs.extend(section.footer.paragraphs)

        for table in section.header.tables:
            all_paragraphs.extend(extract_table_paragraphs(table))

        for table in section.footer.tables:
            all_paragraphs.extend(extract_table_paragraphs(table))
        


    for i, paragraph in enumerate(document.paragraphs):
        fulltext = paragraph.text
        matches = pattern.findall(fulltext)


        matches = pattern.findall(paragraph.text)
        if matches:
            placeholder_text = matches[0]

            label = fulltext.split(placeholder_text)[0].strip()

            placeholders.append({
                "type": "paragraph",
                "index": i,
                "paragraph_obj": paragraph,   
                "text": fulltext,
                "placeholder": placeholder_text,
                "label": label
            })

    return placeholders

